#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX 기획서에서 텍스트, 표, 이미지를 추출하여 공통 스키마로 반환.
Usage: python extract_docx.py <docx_path> [--reference-dir DIR]
"""

import os
import sys
import zipfile

from _utils import parse_reference_dir, resolve_ref_dir, setup_stdout_utf8, to_rel_path


def extract_docx(docx_path, reference_dir=None):
    """DOCX 파일에서 내용 추출. 공통 스키마 반환."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx 라이브러리가 필요합니다.")
        print("설치: pip install python-docx")
        return None

    if not os.path.exists(docx_path):
        print(f"Error: 파일을 찾을 수 없습니다 - {docx_path}")
        return None

    ref_dir = resolve_ref_dir(reference_dir)

    doc = Document(docx_path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = [[[cell.text.strip() for cell in row.cells] for row in t.rows] for t in doc.tables]
    images = []
    reference_images = []

    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            for name in z.namelist():
                if not name.startswith("word/media/"):
                    continue
                ext = os.path.splitext(name)[1].lower()
                if ext not in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
                    continue
                data = z.read(name)
                if len(data) < 500:
                    continue
                img_index = len(images)
                out_ext = ".jpg" if ext in (".jpg", ".jpeg") else ext
                rel_name = f"docx_page_1_img_{img_index}{out_ext}"
                abs_path = os.path.join(ref_dir, rel_name)
                with open(abs_path, "wb") as f:
                    f.write(data)
                rel_path = to_rel_path(abs_path, os.path.join("outputs", "reference", rel_name))
                images.append({"path": rel_path, "description": ""})
                reference_images.append({"source_page": 1, "path": rel_path})
    except Exception as e:
        print(f"Warning: 이미지 추출 중 오류 - {e}")

    pages_out = [{
        "page_num": 1,
        "texts": texts,
        "tables": tables,
        "notes": "",
        "images": images,
    }]
    result = {"pages": pages_out, "reference_images": reference_images}

    print(f"=== {os.path.basename(docx_path)} 분석 결과 ===\n")
    print("--- 페이지 1 ---")
    for t in texts:
        print(t)
    for table in tables:
        for row in table:
            print(" | ".join(row))
        print()
    for im in images:
        if im.get("path"):
            print(f"  [이미지] {im['path']}")
    print()

    return result


def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_path> [--reference-dir DIR]")
        sys.exit(1)

    setup_stdout_utf8()
    r = extract_docx(sys.argv[1], reference_dir=parse_reference_dir(sys.argv))
    sys.exit(0 if r is not None else 1)


if __name__ == "__main__":
    main()
